"""
Document processing for multi-turn context.

Processes documents for LLM consumption:
- PDF: Pass through as native document block (Claude handles natively)
- DOCX/XLSX: Extract text using optional libraries with stdlib fallback

Optional dependencies (python-docx, openpyxl) provide richer extraction
but are not required - falls back to stdlib XML parsing.
"""
import base64
import logging
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Optional, TYPE_CHECKING
from xml.etree import ElementTree

if TYPE_CHECKING:
    from clients.files_manager import FilesManager

logger = logging.getLogger(__name__)

# Try importing optional libraries
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.debug("python-docx not installed - using stdlib fallback for DOCX")

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    logger.debug("openpyxl not installed - using stdlib fallback for XLSX")


SUPPORTED_DOCUMENT_FORMATS = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # XLSX
    "text/plain",  # TXT
    "text/csv",  # CSV
    "application/json",  # JSON
}

MAX_DOCUMENT_SIZE_MB = 32  # Claude's PDF limit


@dataclass(frozen=True)
class ProcessedDocument:
    """
    Result of document processing.

    content_type values:
    - "container_upload": Structured data uploaded to Files API (CSV, XLSX, JSON)
    - "document": PDF as base64 (not code-executable)
    - "text": Extracted text from DOCX
    """

    content_type: str   # "container_upload" | "document" | "text"
    media_type: str     # Original MIME type
    data: str           # file_id | base64 | extracted_text
    original_filename: Optional[str] = None  # Filename for tracking


def process_document(
    doc_bytes: bytes,
    media_type: str,
    files_manager: Optional['FilesManager'] = None,
    filename: str = "document",
    segment_id: Optional[str] = None
) -> ProcessedDocument:
    """
    Process document for LLM consumption with Files API support.

    Type-based routing:
    - CSV/XLSX/JSON → Upload to Files API (code-executable)
    - PDF → Base64 document block (display-only)
    - DOCX → Extract text (Anthropic doesn't support native)

    Args:
        doc_bytes: Raw document bytes
        media_type: MIME type of the document
        files_manager: Optional FilesManager for Files API uploads
        filename: Original filename for Files API tracking

    Returns:
        ProcessedDocument with appropriate content for LLM

    Raises:
        ValueError: If document type unsupported or processing fails
    """
    # Structured data: Upload to Files API for code execution
    if media_type in ("text/csv", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/json"):
        if files_manager is not None and segment_id is not None:
            # Upload to Files API, return file_id
            file_id = files_manager.upload_file(
                file_bytes=doc_bytes,
                filename=filename,
                media_type=media_type,
                segment_id=segment_id
            )
            return ProcessedDocument(
                content_type="container_upload",
                media_type=media_type,
                data=file_id,
                original_filename=filename
            )
        else:
            # Fallback: extract text if no files_manager provided
            if media_type == "text/csv":
                text = extract_text_file(doc_bytes)
            elif media_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                text = extract_xlsx_text(doc_bytes)
            elif media_type == "application/json":
                text = extract_text_file(doc_bytes)
            else:
                text = extract_text_file(doc_bytes)

            return ProcessedDocument(
                content_type="text",
                media_type=media_type,
                data=text
            )

    # PDF: Pass through as base64 document block (not code-executable)
    elif media_type == "application/pdf":
        return ProcessedDocument(
            content_type="document",
            media_type=media_type,
            data=base64.b64encode(doc_bytes).decode('utf-8')
        )

    # DOCX: Extract text (Anthropic doesn't support native DOCX)
    elif media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_docx_text(doc_bytes)
        return ProcessedDocument(
            content_type="text",
            media_type=media_type,
            data=text
        )

    # Plain text: Extract as UTF-8
    elif media_type == "text/plain":
        text = extract_text_file(doc_bytes)
        return ProcessedDocument(
            content_type="text",
            media_type=media_type,
            data=text
        )

    else:
        raise ValueError(f"Unsupported document type: {media_type}")


def extract_text_file(doc_bytes: bytes) -> str:
    """
    Extract text from TXT or CSV file.

    Attempts UTF-8 decoding, falls back to latin-1 if that fails.
    """
    try:
        return doc_bytes.decode('utf-8')
    except UnicodeDecodeError:
        return doc_bytes.decode('latin-1')


def extract_docx_text(doc_bytes: bytes) -> str:
    """
    Extract text from DOCX document.

    Uses python-docx if available, otherwise falls back to stdlib XML parsing.
    """
    if HAS_DOCX:
        return _extract_docx_with_library(doc_bytes)
    else:
        return _extract_docx_stdlib(doc_bytes)


def extract_xlsx_text(doc_bytes: bytes) -> str:
    """
    Extract text from XLSX spreadsheet.

    Uses openpyxl if available, otherwise falls back to stdlib XML parsing.
    """
    if HAS_OPENPYXL:
        return _extract_xlsx_with_library(doc_bytes)
    else:
        return _extract_xlsx_stdlib(doc_bytes)


def _extract_docx_with_library(doc_bytes: bytes) -> str:
    """Extract DOCX text using python-docx library."""
    try:
        doc = docx.Document(BytesIO(doc_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n'.join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {e}") from e


def _extract_docx_stdlib(doc_bytes: bytes) -> str:
    """Extract DOCX text using stdlib only (zipfile + xml)."""
    try:
        with zipfile.ZipFile(BytesIO(doc_bytes)) as z:
            xml_content = z.read('word/document.xml')
            tree = ElementTree.fromstring(xml_content)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            texts = [t.text for t in tree.findall('.//w:t', ns) if t.text]
            return ' '.join(texts)
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {e}") from e


def _extract_xlsx_with_library(doc_bytes: bytes) -> str:
    """Extract XLSX text using openpyxl library."""
    try:
        wb = openpyxl.load_workbook(BytesIO(doc_bytes), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(cells):
                    lines.append('\t'.join(cells))
        wb.close()
        return '\n'.join(lines)
    except Exception as e:
        raise ValueError(f"Failed to extract XLSX text: {e}") from e


def _extract_xlsx_stdlib(doc_bytes: bytes) -> str:
    """Extract XLSX text using stdlib only (zipfile + xml)."""
    try:
        with zipfile.ZipFile(BytesIO(doc_bytes)) as z:
            # Read shared strings (XLSX stores text in a shared strings table)
            shared_strings = []
            if 'xl/sharedStrings.xml' in z.namelist():
                ss_xml = z.read('xl/sharedStrings.xml')
                ss_tree = ElementTree.fromstring(ss_xml)
                # Find all <t> elements (text content)
                for si in ss_tree.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t'):
                    shared_strings.append(si.text or '')

            # Read sheet data
            lines = []
            for name in sorted(z.namelist()):
                if name.startswith('xl/worksheets/sheet') and name.endswith('.xml'):
                    sheet_xml = z.read(name)
                    sheet_tree = ElementTree.fromstring(sheet_xml)
                    ns = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

                    for row in sheet_tree.findall(f'.//{{{ns}}}row'):
                        cells = []
                        for cell in row.findall(f'{{{ns}}}c'):
                            val = cell.find(f'{{{ns}}}v')
                            cell_type = cell.get('t')
                            if val is not None and val.text:
                                if cell_type == 's':  # Shared string reference
                                    idx = int(val.text)
                                    cells.append(shared_strings[idx] if idx < len(shared_strings) else '')
                                else:
                                    cells.append(val.text)
                            else:
                                cells.append('')
                        if any(cells):
                            lines.append('\t'.join(cells))

            return '\n'.join(lines)
    except Exception as e:
        raise ValueError(f"Failed to extract XLSX text: {e}") from e
